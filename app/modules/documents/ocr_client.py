"""
AI-Sourcing Hub — Document Extraction Pipeline

Two-phase pipeline:

Phase 0 — Structured extraction (no LLM):
    • Excel / CSV / Word tables → pandas direct read (header row auto-detected)
    • PDF with drawn tables → pdfplumber (rows/columns preserved)

Phase 1 — Raw text extraction (when Phase 0 finds nothing):
    • PDF  → pypdfium2 (instant, no ML, handles text-based PDFs)
    • Image / scanned PDF → PaddleOCR (skipped gracefully if runtime broken)

Phase 2 — LLM structuring:
    • Raw text is sent to the LLM with a dynamic extraction prompt, split into
      chunks so long catalogs are never truncated.
    • The LLM reads whatever columns/headers exist in this specific document
      and returns a JSON array where each object has the fields found in that
      document (not a hardcoded set of columns).
    • Falls back to empty list if no API key is configured or all calls fail.
"""
from __future__ import annotations

import asyncio
import json
import os
import re
import threading
from typing import Optional

from app.shared.logging import get_logger

logger = get_logger(__name__)

MAX_PDF_PAGES = 20

# ═══════════════════════════════════════════════════════════
# Phase 1A — pypdfium2 text extraction (PDFs)
# ═══════════════════════════════════════════════════════════


def _pdf_to_text(file_bytes: bytes) -> str:
    """Extract all text from a PDF using pypdfium2.

    Returns the raw multi-line text string, or '' if extraction fails.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError:
        logger.warning("pypdfium2 not installed — skipping PDF text extraction")
        return ""

    try:
        pdf = pdfium.PdfDocument(file_bytes)
    except Exception as exc:
        logger.warning("pypdfium2 could not open PDF: %s", exc)
        return ""

    pages_text: list[str] = []
    for page_idx in range(min(len(pdf), MAX_PDF_PAGES)):
        try:
            page = pdf[page_idx]
            text_page = page.get_textpage()
            try:
                text = "".join(text_page.get_text_bounded())
            except (AttributeError, TypeError):
                text = text_page.get_text_range()
            if text.strip():
                pages_text.append(text.strip())
        except Exception as exc:
            logger.debug("pypdfium2 page %d error: %s", page_idx, exc)

    return "\n\n".join(pages_text)


# ═══════════════════════════════════════════════════════════
# Phase 0 (PDF) — pdfplumber structured table extraction
#
# pypdfium2 only gives flat text with no column info, so a 6-column table
# turns into scrambled lines and the LLM has to guess which number belongs
# to which column (prices get mixed with quantities). pdfplumber detects
# actual table structure (rows + columns) from the PDF's drawing geometry,
# so when it succeeds the result feeds straight into the same no-LLM
# row→product pipeline used for Excel/Word.
# ═══════════════════════════════════════════════════════════


def _pdf_to_tables(file_bytes: bytes, max_pages: int) -> list:
    """Extract structured tables from a text-based PDF via pdfplumber.

    Tables that continue across pages usually repeat on each page without
    their header row, so consecutive tables with the same column count are
    concatenated before header detection (which also drops repeated header
    rows). Returns a list of clean pandas DataFrames, or [] if pdfplumber
    is unavailable or finds nothing.
    """
    import io

    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed — skipping PDF table extraction")
        return []

    import pandas as pd

    raw_tables: list[list[list]] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages[:max_pages]:
                for table in page.extract_tables():
                    rows = [r for r in (table or []) if any(
                        c is not None and str(c).strip() for c in r
                    )]
                    if rows and max(len(r) for r in rows) >= 2:
                        raw_tables.append(rows)
    except Exception as exc:
        logger.warning("pdfplumber table extraction failed: %s", exc)
        return []

    # Merge consecutive same-width tables (multi-page continuation).
    groups: list[list[list]] = []
    for table in raw_tables:
        width = max(len(r) for r in table)
        padded = [list(r) + [None] * (width - len(r)) for r in table]
        if groups and len(groups[-1][0]) == width:
            groups[-1].extend(padded)
        else:
            groups.append(padded)

    frames = []
    for rows in groups:
        if len(rows) < 2:
            continue  # header only — no data
        df = _rebuild_with_detected_header(pd.DataFrame(rows, dtype=object))
        if df is not None and not df.empty:
            frames.append(df)
    return frames


# ═══════════════════════════════════════════════════════════
# Phase 1B — PaddleOCR (images / scanned PDFs)
# ═══════════════════════════════════════════════════════════

_ocr_engine = None
_ocr_lock = threading.Lock()
_ocr_unavailable = False  # set True permanently on first runtime failure


def _get_ocr_engine():
    global _ocr_engine, _ocr_unavailable
    if _ocr_unavailable:
        return None
    if _ocr_engine is not None:
        return _ocr_engine
    with _ocr_lock:
        if _ocr_engine is not None:
            return _ocr_engine
        if _ocr_unavailable:
            return None
        try:
            from paddleocr import PaddleOCR

            from app.config import settings
            _ocr_engine = PaddleOCR(
                lang=getattr(settings, "OCR_LANG", "en"),
                use_textline_orientation=False,
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
            )
            logger.info("PaddleOCR engine ready (lang=%s)", getattr(settings, "OCR_LANG", "en"))
        except Exception as exc:
            logger.warning("PaddleOCR unavailable (%s) — will skip OCR step", exc)
            _ocr_unavailable = True
    return _ocr_engine


def warm_up_ocr() -> None:
    """Force-load the PaddleOCR engine so the first real request isn't slow.

    Safe to call from app startup — errors are swallowed (mirrors the lazy
    lookup's own error handling), the OCR feature is simply unavailable.
    """
    try:
        _get_ocr_engine()
    except Exception:
        pass


def _ocr_image_to_text(img_path: str) -> str:
    """Run PaddleOCR on an image and return extracted text as a string."""
    engine = _get_ocr_engine()
    if engine is None:
        return ""

    try:
        results = engine.predict(
            img_path,
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
        )
    except Exception as exc:
        # A single bad image shouldn't permanently disable OCR for the
        # process — only _get_ocr_engine() failing to load the model does.
        logger.warning("PaddleOCR predict failed: %s", exc)
        return ""

    # Collect (y_center, x_left, text) then sort top→bottom, left→right
    boxes: list[tuple[float, float, str]] = []
    for page_result in (results or []):
        try:
            rec_texts = getattr(page_result, "rec_texts", None) or []
            # PaddleOCR 3.x pairs rec_texts with rec_boxes/rec_polys (index
            # aligned); dt_polys is the pre-recognition detection set and can
            # have a different length, so it must not be used here.
            rec_boxes = getattr(page_result, "rec_boxes", None)
            rec_polys = getattr(page_result, "rec_polys", None) or []
            if rec_texts and rec_boxes is not None and len(rec_boxes):
                for box, text in zip(rec_boxes, rec_texts):
                    x1, y1, x2, y2 = box[0], box[1], box[2], box[3]
                    boxes.append(((y1 + y2) / 2, min(x1, x2), text))
                continue
            if rec_texts and rec_polys:
                for box, text in zip(rec_polys, rec_texts):
                    ys = [p[1] for p in box]; xs = [p[0] for p in box]
                    boxes.append((sum(ys)/len(ys), min(xs), text))
                continue
        except Exception:
            pass
        # Paddle 2.x list-of-tuple fallback
        try:
            for item in page_result:
                box, (text, _) = item
                ys = [p[1] for p in box]; xs = [p[0] for p in box]
                boxes.append((sum(ys)/len(ys), min(xs), text))
        except Exception:
            pass

    if not boxes:
        return ""

    boxes.sort(key=lambda t: t[0])
    lines: list[str] = []
    current_row: list[tuple[float, str]] = []
    y0 = boxes[0][0]
    Y_THRESHOLD = 15.0

    for y, x, text in boxes:
        if abs(y - y0) <= Y_THRESHOLD:
            current_row.append((x, text))
        else:
            if current_row:
                lines.append("  ".join(t for _, t in sorted(current_row)))
            current_row = [(x, text)]
            y0 = y
    if current_row:
        lines.append("  ".join(t for _, t in sorted(current_row)))

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════
# Shared field/header map (used by BOTH the LLM prompt and the
# direct table readers — one source of truth, no duplication)
# ═══════════════════════════════════════════════════════════

# Each entry: list of source labels (Chinese / Arabic / English) → English field.
# The LLM prompt's "common mappings" bullet list is generated from this, and the
# direct Excel/CSV/Word table readers translate their column headers through it.
_FIELD_ALIASES: list[tuple[list[str], str]] = [
    (["产品名称", "اسم المنتج", "Item Name", "Product Name"], "product_name"),
    (["型号", "رقم الموديل", "Model", "SKU", "Article"], "model_number"),
    (["颜色", "Color", "Colour", "اللون"], "color"),
    (["尺寸", "Size", "Dimension", "المقاس"], "size"),
    (["材质", "Material", "الخامة", "الخامات"], "material"),
    (["重量", "Weight", "الوزن"], "weight"),
    (["起订量", "MOQ", "الحد الأدنى"], "moq"),
    (["出厂单价", "Unit Price", "السعر"], "unit_price"),
    (["سعر الجملة", "Wholesale Price"], "wholesale_price"),
    (["سعر التجزئة", "Retail Price"], "retail_price"),
    (["الكمية", "Quantity", "Stock"], "quantity"),
    (["الكود", "Code", "Ref", "Reference"], "code"),
    (["الوصف", "Description", "Notes"], "description"),
    (["الفئة", "Category", "Type", "نوع"], "category"),
    (["العلامة التجارية", "Brand", "ماركة"], "brand"),
    (["Origin", "Country of Origin", "بلد المنشأ"], "origin"),
    (["Voltage"], "voltage"),
    (["Wattage"], "wattage"),
    (["Power"], "power"),
    (["Expiry", "Shelf Life", "تاريخ الانتهاء"], "shelf_life"),
    (["Fabric", "Composition", "التركيب"], "fabric_composition"),
    (["Season", "الموسم"], "season"),
    (["Gender", "الجنس"], "gender"),
    (["Packaging", "التعبئة"], "packaging"),
    (["Certification", "شهادة"], "certification"),
]

# Normalised (stripped + lowercased) source label → English field.
_HEADER_MAP: dict[str, str] = {
    alias.strip().lower(): field
    for aliases, field in _FIELD_ALIASES
    for alias in aliases
}

# Human-readable bullet list injected into the LLM system prompt.
_MAPPING_TEXT = "\n".join(
    "    " + " / ".join(aliases) + " → " + field for aliases, field in _FIELD_ALIASES
)


def _snake_case(header: str) -> str:
    """Normalise an arbitrary column header to a snake_case field name.

    Strips, lowercases, turns any run of non-alphanumeric chars into a single
    underscore, and trims leading/trailing underscores. Non-ASCII letters are
    kept (so untranslated Arabic/Chinese headers still yield a stable key).
    """
    s = header.strip().lower()
    s = re.sub(r"[^0-9a-z؀-ۿ一-鿿]+", "_", s)
    return s.strip("_")


def _translate_header(header: str) -> str:
    """Map a raw column header to its English field name via _HEADER_MAP,
    falling back to a snake_cased version of the header itself."""
    return _HEADER_MAP.get(header.strip().lower()) or _snake_case(header)


# ═══════════════════════════════════════════════════════════
# Phase 2 — LLM structuring
# ═══════════════════════════════════════════════════════════

_SYSTEM_PROMPT = (
    "You are a universal product data extractor for B2B supplier catalogs.\n"
    "You work with ANY type of catalog: electronics, furniture, textiles, clothing, "
    "food, chemicals, machinery, lighting, construction materials, cosmetics, medical "
    "supplies, or anything else. Never assume what industry you are in — read the "
    "document and adapt.\n\n"
    "Your job:\n"
    "1. Read the supplier document text provided by the user.\n"
    "2. Identify ALL distinct products or items in the document.\n"
    "3. For each product, extract EVERY field that appears in the document for that product.\n"
    "   - Use the document's own column headers / labels as field names (translate to English).\n"
    "   - Do NOT map to a fixed schema — every catalog is different.\n"
    "4. Return ONLY a valid JSON array. No markdown fences, no explanation, nothing else.\n\n"
    "Field naming rules:\n"
    '- Translate any non-English field name to English (snake_case, e.g. "unit_price").\n'
    '- Always include "product_name" (or the closest equivalent in the document).\n'
    "- Common label mappings (not exhaustive — use whatever the document actually has):\n"
    + _MAPPING_TEXT
    + "\n\n"
    'Numeric values: store as numbers, not strings (e.g. 12.5 not "12.5").\n'
    "Missing field for a product: omit the key entirely — do NOT use null.\n"
    'If the document has sections or categories, add a "category" field to each product.\n\n'
    "Output format (example — your actual fields will differ):\n"
    "[\n"
    '  {"product_name": "...", "model_number": "...", "unit_price": 12.5, "moq": 500, "color": "red", ...},\n'
    '  {"product_name": "...", ...}\n'
    "]\n"
)


# Chunk size chosen to stay well inside free-tier context limits while
# leaving room for the system prompt and an 8K-token JSON response.
_CHUNK_CHARS = 6000
_MAX_PARALLEL_LLM_CALLS = 3


def _split_text_chunks(raw_text: str, chunk_chars: int = _CHUNK_CHARS) -> list[str]:
    """Split text into ~chunk_chars pieces on line boundaries, so a product's
    row is never cut in half between two chunks."""
    chunks: list[str] = []
    buf: list[str] = []
    buf_len = 0
    for line in raw_text.split("\n"):
        if buf and buf_len + len(line) + 1 > chunk_chars:
            chunks.append("\n".join(buf))
            buf, buf_len = [], 0
        buf.append(line)
        buf_len += len(line) + 1
    if buf:
        chunks.append("\n".join(buf))
    return chunks


async def _llm_extract(raw_text: str) -> list[dict]:
    """Send raw document text to the LLM and return the structured product
    list. Long documents are split into chunks (all products get extracted,
    nothing is truncated away) processed concurrently with a small cap to
    stay under provider rate limits."""
    if len(raw_text) <= _CHUNK_CHARS:
        return await _llm_extract_chunk(raw_text)

    chunks = _split_text_chunks(raw_text)
    logger.info(
        "Document text split into %d chunks (%d chars total) for LLM extraction",
        len(chunks), len(raw_text),
    )

    sem = asyncio.Semaphore(_MAX_PARALLEL_LLM_CALLS)

    async def _run(chunk: str) -> list[dict]:
        async with sem:
            return await _llm_extract_chunk(chunk)

    results = await asyncio.gather(*(_run(c) for c in chunks))

    products: list[dict] = []
    for i, chunk_products in enumerate(results):
        if chunk_products:
            products.extend(chunk_products)
        else:
            logger.warning("LLM chunk %d/%d returned no products", i + 1, len(chunks))
    return products


async def _llm_extract_chunk(raw_text: str) -> list[dict]:
    """Send one chunk of document text to the LLM and return structured
    products. Tries DeepSeek, then OpenRouter models, then Together AI.
    Returns [] if unavailable.
    """
    try:
        from app.config import settings
    except Exception:
        return []

    deepseek_key   = getattr(settings, "DEEPSEEK_API_KEY",   "") or ""
    openrouter_key = getattr(settings, "OPENROUTER_API_KEY", "") or ""
    together_key   = getattr(settings, "TOGETHER_API_KEY",   "") or ""

    if not deepseek_key and not openrouter_key and not together_key:
        logger.warning("No LLM API key configured — skipping LLM extraction")
        return []

    # raw_text arrives pre-chunked to ~_CHUNK_CHARS by _llm_extract, so no
    # truncation happens here — every product gets a chance to be extracted.
    messages = [
        {"role": "system", "content": _SYSTEM_PROMPT},
        {"role": "user",   "content": f"Extract all products from this document:\n\n{raw_text}"},
    ]

    # Free OpenRouter models — tried in order (verified available as of 2026-06).
    # On rate limit (429), immediately falls through to the next model.
    OPENROUTER_MODELS = [
        "meta-llama/llama-3.3-70b-instruct:free",       # primary
        "google/gemma-4-31b-it:free",                    # Google Gemma 4
        "qwen/qwen3-next-80b-a3b-instruct:free",         # Qwen3 — good for Chinese
        "openai/gpt-oss-120b:free",                      # OpenAI OSS
        "openai/gpt-oss-20b:free",                       # smaller fallback
        "meta-llama/llama-3.2-3b-instruct:free",         # tiny but works
    ]

    endpoints = []
    # DeepSeek first — the only provider with a live key (mirrors intake/llm_client.py).
    if deepseek_key:
        endpoints.append({
            "url":   "https://api.deepseek.com/v1/chat/completions",
            "key":   deepseek_key,
            "model": "deepseek-chat",
        })
    if openrouter_key:
        for model in OPENROUTER_MODELS:
            endpoints.append({
                "url":   "https://openrouter.ai/api/v1/chat/completions",
                "key":   openrouter_key,
                "model": model,
            })
    if together_key:
        endpoints.append({
            "url":   "https://api.together.xyz/v1/chat/completions",
            "key":   together_key,
            "model": "meta-llama/Llama-3.3-70B-Instruct-Turbo",
        })

    import httpx

    for ep in endpoints:
        try:
            async with httpx.AsyncClient(timeout=20.0) as client:
                resp = await client.post(
                    ep["url"],
                    headers={
                        "Authorization": f"Bearer {ep['key']}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model":       ep["model"],
                        "messages":    messages,
                        "temperature": 0.1,
                        "max_tokens":  8192,
                    },
                )

            if resp.status_code == 429:
                logger.warning(
                    "Rate limited by %s model=%s — trying next model",
                    ep["url"], ep["model"],
                )
                continue  # next model immediately, no sleep

            if resp.status_code in (401, 403):
                logger.warning("Auth error for %s model=%s — skipping", ep["url"], ep["model"])
                continue

            resp.raise_for_status()
            content = resp.json()["choices"][0]["message"]["content"].strip()

            # Strip markdown fences if the model wrapped the JSON
            content = re.sub(r"^```[a-z]*\s*", "", content, flags=re.IGNORECASE)
            content = re.sub(r"\s*```$",        "", content)
            content = content.strip()

            # Find the JSON array (handle preamble text from chatty models)
            array_start = content.find("[")
            if array_start != -1:
                content = content[array_start:]

            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                # Usually the array got cut off mid-object at max_tokens.
                # json-repair closes unterminated structures, recovering the
                # complete products instead of discarding the whole batch.
                from json_repair import repair_json

                data = json.loads(repair_json(content))
                if isinstance(data, list) and data:
                    # The truncated tail may come back as a partial dict or
                    # even a stray list — drop it, keep complete dicts only.
                    tail = data[-1]
                    if not isinstance(tail, dict) or "product_name" not in tail:
                        data = data[:-1]
                    data = [d for d in data if isinstance(d, dict)]
                logger.warning(
                    "JSON from model=%s was truncated — recovered %d complete products",
                    ep["model"], len(data) if isinstance(data, list) else 0,
                )

            if isinstance(data, dict) and "products" in data:
                data = data["products"]

            if isinstance(data, list) and data:
                logger.info(
                    "LLM extracted %d products — model=%s", len(data), ep["model"]
                )
                return data

            logger.warning("LLM returned empty list — model=%s", ep["model"])

        except json.JSONDecodeError as exc:
            logger.warning("Invalid JSON from model=%s: %s", ep["model"], exc)
        except Exception as exc:
            logger.warning("LLM call error model=%s: %s", ep["model"], exc)

    return []


# ═══════════════════════════════════════════════════════════
# Direct structured-file readers (NO LLM) — Excel / CSV / Word tables
#
# Structured inputs already ARE tables: one row = one product, one column =
# one field. We read them directly for near-exact accuracy and instant speed,
# and only fall back to the LLM when a table is too ambiguous to trust.
# ═══════════════════════════════════════════════════════════

IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif")


def _coerce_number(text: str):
    """Convert a numeric-looking string to int/float, else return None.

    Leading-zero integer strings (e.g. "007", "0123") are treated as codes and
    left as strings, so SKUs/reference numbers are not silently mangled.
    """
    t = text.strip()
    if re.fullmatch(r"-?\d+", t):
        digits = t.lstrip("-")
        if len(digits) > 1 and digits.startswith("0"):
            return None  # leading-zero code — keep as string
        try:
            return int(t)
        except ValueError:
            return None
    if re.fullmatch(r"-?(\d+\.\d*|\.\d+)", t):
        try:
            return float(t)
        except ValueError:
            return None
    return None


def _clean_cell(value):
    """Return a JSON-safe scalar for a cell, or None if it is empty/NaN.

    Numeric values are returned as native int/float (not numpy scalars, not
    strings), matching the LLM path's "numbers as numbers" contract. Text-based
    formats (CSV, Word) store numbers as strings, so numeric-looking strings are
    coerced too.
    """
    # pandas / numpy NaN and None
    try:
        import pandas as pd

        if value is None or (not isinstance(value, (list, dict)) and pd.isna(value)):
            return None
    except Exception:
        if value is None:
            return None

    # Native-ise numpy scalars via .item()
    item = getattr(value, "item", None)
    if callable(item):
        try:
            value = value.item()
        except Exception:
            pass

    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return None
        number = _coerce_number(stripped)
        return number if number is not None else stripped

    if isinstance(value, bool):
        return value

    if isinstance(value, (int, float)):
        # Collapse integral floats (e.g. 500.0 → 500) for cleaner output
        if isinstance(value, float) and value.is_integer():
            return int(value)
        return value

    text = str(value).strip()
    return text or None


def _rows_to_products(df) -> list[dict]:
    """Convert a DataFrame (first row already used as header) into product dicts.

    - Column headers are translated to English snake_case via _HEADER_MAP.
    - Empty/NaN cells are omitted (never null).
    - Numeric cells become real numbers.
    - Fully-empty rows are skipped.
    """
    columns = [_translate_header(str(c)) for c in df.columns]

    products: list[dict] = []
    for _, row in df.iterrows():
        product: dict = {}
        for field, raw in zip(columns, row.tolist()):
            cleaned = _clean_cell(raw)
            if cleaned is not None and field:
                product[field] = cleaned
        if product:
            products.append(product)
    return products


def _is_ambiguous(df) -> bool:
    """A table is too ambiguous to read directly if it has fewer than two
    usable columns, or nearly every header is a placeholder — pandas
    'Unnamed: N' or our own 'col_N' from header detection."""
    cols = [str(c).strip() for c in df.columns]
    if len(cols) < 2:
        return True
    meaningful = [
        c for c in cols
        if c
        and not c.lower().startswith("unnamed")
        and not re.fullmatch(r"col_\d+(_\d+)?", c.lower())
    ]
    return len(meaningful) < 2


# ── Header-row detection ────────────────────────────────────
# Real supplier catalogs rarely start with the header row: there is usually
# a merged title row, logo/contact/date rows, and THEN the actual header.
# Reading with header=0 turns that title row into the header ("Unnamed: N"
# columns), which either mistranslates every field or needlessly dumps the
# whole file onto the LLM. So tables are read headerless and the most
# header-looking row within the first rows is detected and promoted.


def _cell_text(value) -> str:
    """Raw cell → stripped string, '' for None/NaN."""
    if value is None:
        return ""
    if isinstance(value, float) and value != value:  # NaN
        return ""
    return str(value).strip()


def _row_header_score(row: list) -> float:
    """Score how much a raw row looks like a table header row.

    Higher = more short non-numeric text cells and/or cells matching a known
    field alias. Merged-title rows score low (few filled cells); data rows
    score low (mostly numeric cells).
    """
    cells = [c for c in (_cell_text(v) for v in row) if c]
    if not cells:
        return float("-inf")
    score = 0.0
    for c in cells:
        if c.lower() in _HEADER_MAP:
            score += 3.0  # strong signal: matches a known field alias
        elif _coerce_number(c) is not None:
            score -= 1.0  # numeric cell — looks like data, not a header
        elif len(c) <= 40:
            score += 1.0  # short text cell — plausible header label
    # Reward more distinct filled cells (a title row usually has just one).
    score += 0.3 * len(cells)
    return score


def _detect_header_row(raw_rows: list[list], max_scan: int = 10) -> int:
    """Return the index of the most header-looking row among the first
    `max_scan` rows."""
    best_idx, best_score = 0, float("-inf")
    for idx, row in enumerate(raw_rows[:max_scan]):
        score = _row_header_score(row)
        if score > best_score:
            best_idx, best_score = idx, score
    return best_idx


def _rebuild_with_detected_header(df_raw):
    """Turn a headerless DataFrame into a clean one with proper columns.

    Detects the real header row, deduplicates/fills merged header cells,
    drops repeated header rows (multi-page PDF tables repeat theirs on every
    page), and forward-fills vertically-merged cells — but only in mostly-
    text columns, so a price is never silently duplicated into blank rows.
    Returns None when there is no data below the detected header.
    """
    raw_rows = df_raw.values.tolist()
    if not raw_rows:
        return None

    header_idx = _detect_header_row(raw_rows)
    header_cells = df_raw.iloc[header_idx].ffill().tolist()

    names: list[str] = []
    seen: dict[str, int] = {}
    for i, h in enumerate(header_cells):
        name = _cell_text(h) or f"col_{i}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        names.append(name)

    body = df_raw.iloc[header_idx + 1:].copy()
    if body.empty:
        return None
    body.columns = names

    header_set = {n.lower() for n in names}

    def _is_header_repeat(row) -> bool:
        vals = [c.lower() for c in (_cell_text(v) for v in row.tolist()) if c]
        return bool(vals) and sum(v in header_set for v in vals) > len(vals) / 2

    body = body[~body.apply(_is_header_repeat, axis=1)]

    # Positional access: header ffill can produce duplicate names, and
    # body[name] would then return a DataFrame instead of a Series.
    for i in range(len(names)):
        series = body.iloc[:, i]
        non_null = series.dropna()
        if non_null.empty:
            continue
        numeric_like = sum(_coerce_number(str(v)) is not None for v in non_null)
        if numeric_like / len(non_null) < 0.5:
            body.iloc[:, i] = series.ffill()

    return body.reset_index(drop=True)


async def _table_file_to_products(file_bytes: bytes, ext: str) -> list[dict]:
    """Read an Excel / CSV / TSV file directly into product dicts (no LLM).

    Files are read headerless first, then the real header row is auto-detected
    (skipping the merged title/logo/contact rows common in supplier catalogs)
    and vertically-merged cells are filled. Falls back to the LLM only when
    the resulting table is still too ambiguous to trust.
    """
    import io

    import pandas as pd

    frames: list = []
    try:
        if ext in (".xlsx", ".xls"):
            engine = "openpyxl" if ext == ".xlsx" else None
            sheets = pd.read_excel(
                io.BytesIO(file_bytes), sheet_name=None, engine=engine,
                dtype=object, header=None,
            )
            raw_frames = [df for df in sheets.values() if not df.empty]
        else:  # .csv / .tsv
            sep = "\t" if ext == ".tsv" else ","
            raw_frames = [
                pd.read_csv(io.BytesIO(file_bytes), sep=sep, dtype=object,
                            skip_blank_lines=True, header=None)
            ]
        frames = [_rebuild_with_detected_header(df) for df in raw_frames]
    except Exception as exc:
        logger.warning("Failed to parse %s table (%s) — falling back to LLM", ext, exc)
        frames = []

    frames = [df for df in frames if df is not None and not df.empty]

    if not frames:
        return []

    # If any sheet is a clean table, read it directly; collect ambiguous ones for LLM fallback.
    products: list[dict] = []
    ambiguous_frames: list = []
    for df in frames:
        if _is_ambiguous(df):
            ambiguous_frames.append(df)
        else:
            products.extend(_rows_to_products(df))

    if products:
        logger.info("Read %d products directly from %s table (no LLM)", len(products), ext)
        return products

    # Nothing clean — last resort: dump the raw table text to the LLM.
    logger.info("%s table looks ambiguous — falling back to LLM structuring", ext)
    raw_text = "\n\n".join(df.to_csv(index=False) for df in (ambiguous_frames or frames))
    return await _llm_extract(raw_text)


async def _docx_to_products(file_bytes: bytes) -> list[dict]:
    """Read a .docx file: tables go through the direct reader (no LLM);
    prose-only documents fall back to the LLM path."""
    import io

    try:
        import docx  # python-docx
    except ImportError:
        logger.warning("python-docx not installed — skipping .docx extraction")
        return []

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.warning("python-docx could not open document: %s", exc)
        return []

    # ── Tables → direct read (no LLM) ──────────────────────────
    if document.tables:
        import pandas as pd

        products: list[dict] = []
        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue  # header only or empty
            df = pd.DataFrame(rows[1:], columns=rows[0])
            if not _is_ambiguous(df):
                products.extend(_rows_to_products(df))
        if products:
            logger.info("Read %d products directly from .docx tables (no LLM)", len(products))
            return products

    # ── Prose only (or unreadable tables) → LLM ────────────────
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if not text.strip():
        logger.warning(".docx has no readable text or tables — returning empty list")
        return []
    logger.info(".docx has no clean table — falling back to LLM structuring")
    return await _llm_extract(text)


# ═══════════════════════════════════════════════════════════
# Messy-input path (PDF / image)
# ═══════════════════════════════════════════════════════════


async def _extract_pdf_or_image(
    file_bytes: bytes,
    file_name: str,
    ext: str,
    max_pages: int,
) -> list[dict]:
    """Pipeline for messy inputs. PDFs first try structured table extraction
    (pdfplumber, no LLM); otherwise raw text is extracted (pypdfium2 for
    PDFs, PaddleOCR for images/scanned PDFs) and structured by the LLM."""
    import tempfile

    is_pdf = ext == ".pdf"
    raw_text = ""

    # ── Phase 0 (PDF only): structured tables, no LLM ──────────
    # When the PDF has real drawn tables this reads them with correct
    # row/column structure — same accuracy as the Excel/Word direct paths.
    if is_pdf:
        table_frames = await asyncio.get_event_loop().run_in_executor(
            None, _pdf_to_tables, file_bytes, max_pages
        )
        table_products: list[dict] = []
        for df in table_frames:
            if not _is_ambiguous(df):
                table_products.extend(_rows_to_products(df))
        if table_products:
            logger.info(
                "Read %d products directly from PDF tables via pdfplumber (no LLM): %s",
                len(table_products), file_name,
            )
            return table_products

    # ── Phase 1: Extract raw text ──────────────────────────────
    if is_pdf:
        raw_text = await asyncio.get_event_loop().run_in_executor(
            None, _pdf_to_text, file_bytes
        )
        if raw_text.strip():
            logger.info(
                "pypdfium2 extracted %d chars from %s", len(raw_text), file_name
            )
        else:
            logger.info(
                "PDF has no embedded text (%s) — trying PaddleOCR on page images",
                file_name,
            )

    # For images OR scanned PDFs with no embedded text → use PaddleOCR
    if not raw_text.strip():
        image_paths: list[str] = []
        try:
            if is_pdf:
                try:
                    from pdf2image import convert_from_bytes
                    pages = convert_from_bytes(
                        file_bytes, dpi=200, first_page=1, last_page=max_pages
                    )
                    for i, page in enumerate(pages):
                        tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
                        page.save(tmp.name, "JPEG")
                        image_paths.append(tmp.name)
                except Exception as exc:
                    logger.warning("pdf2image failed: %s", exc)
            elif ext in IMAGE_EXTS:
                tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
                tmp.write(file_bytes)
                tmp.close()
                image_paths.append(tmp.name)

            ocr_parts: list[str] = []
            for img_path in image_paths:
                page_text = await asyncio.get_event_loop().run_in_executor(
                    None, _ocr_image_to_text, img_path
                )
                if page_text.strip():
                    ocr_parts.append(page_text.strip())

            raw_text = "\n\n".join(ocr_parts)
            if raw_text.strip():
                logger.info(
                    "PaddleOCR extracted %d chars from %s", len(raw_text), file_name
                )

        finally:
            for p in image_paths:
                try:
                    os.unlink(p)
                except OSError:
                    pass

    if not raw_text.strip():
        logger.warning(
            "Could not extract any text from %s — returning empty product list",
            file_name,
        )
        return []

    # ── Phase 2: LLM structuring ───────────────────────────────
    products = await _llm_extract(raw_text)

    if not products:
        logger.warning(
            "LLM returned no products for %s — document marked extracted with 0 items",
            file_name,
        )

    return products


# ═══════════════════════════════════════════════════════════
# Public async interface (called by Celery task)
# ═══════════════════════════════════════════════════════════


SUPPORTED_EXTS = (
    ".pdf", ".xlsx", ".xls", ".csv", ".tsv", ".docx", ".txt", *IMAGE_EXTS
)


async def extract_table_local(
    file_bytes: bytes,
    file_name: str,
    max_pages: int = MAX_PDF_PAGES,
) -> list[dict]:
    """Extract structured product data from a document file.

    Dispatches by file extension to the most appropriate reader:
        • .pdf / images        → raw text (pypdfium2 / PaddleOCR) → LLM
        • .xlsx/.xls/.csv/.tsv  → direct table read (NO LLM; LLM only as fallback)
        • .docx                 → direct read of tables, else prose → LLM
        • .doc  (old binary)    → not supported yet (convert to .docx)
        • .txt                  → text → LLM directly
        • anything else         → [] with a clear log message

    In every case the return shape is the same: a list of product dicts with
    whatever fields were found (English snake_case keys, numbers as numbers,
    missing fields omitted).

    Args:
        file_bytes: Raw file bytes.
        file_name:  Original file name (used to detect type).
        max_pages:  Max PDF pages to read (default 20).

    Returns:
        List of product dicts.
    """
    ext = os.path.splitext(file_name)[1].lower()

    # ── Structured tables → direct read (no LLM) ───────────────
    if ext in (".xlsx", ".xls", ".csv", ".tsv"):
        return await _table_file_to_products(file_bytes, ext)

    # ── Word ───────────────────────────────────────────────────
    if ext == ".docx":
        return await _docx_to_products(file_bytes)

    if ext == ".doc":
        logger.warning(
            "Old binary .doc not supported (%s) — please convert to .docx", file_name
        )
        return []

    # ── Plain text → straight to LLM (skip Phase 1) ────────────
    if ext == ".txt":
        text = file_bytes.decode("utf-8", errors="replace")
        if not text.strip():
            logger.warning("%s is empty — returning empty product list", file_name)
            return []
        return await _llm_extract(text)

    # ── Messy inputs (PDF / images) → raw text → LLM ───────────
    if ext == ".pdf" or ext in IMAGE_EXTS:
        return await _extract_pdf_or_image(file_bytes, file_name, ext, max_pages)

    logger.warning(
        "Unsupported format '%s' for %s — returning empty list. Supported: %s",
        ext or "(none)",
        file_name,
        ", ".join(SUPPORTED_EXTS),
    )
    return []
